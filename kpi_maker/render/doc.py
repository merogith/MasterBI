"""Editable report (DOCX).

Same eight sections as the PDF. The difference is intent: the PDF is the record,
the DOCX is what a client edits before circulating it internally. So this one
uses real Word heading styles (so the navigation pane and any generated table of
contents work) rather than hand-styled text.
"""
from __future__ import annotations

import io
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from ..fmt import fmt_value
from ..insight.detectors import Finding
from ..kpi.schema import KPISet, Perspective
from ..metrics.engine import MetricResult
from ..profile.schema import CompanyProfile
from ..viz.theme import TOKENS

STATUS_WORD = {"green": "On track", "amber": "Watch", "red": "Off track",
               "unscored": "No target", "unknown": "No data"}
SEVERITY_WORD = {"critical": "Critical", "high": "High", "medium": "Medium",
                 "low": "Low", "positive": "Strength"}


def _rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    for name, size, color in (("Heading 1", 18, "text_primary"),
                              ("Heading 2", 13, "text_primary"),
                              ("Heading 3", 11, "text_primary")):
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        st.font.name = "Segoe UI"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = _rgb(doc.t[color])


def _muted(doc: Document, text: str, size: int = 9, italic: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.font.color.rgb = _rgb(doc.t["muted"])
    return p


def _table(doc: Document, headers: List[str], rows: List[List[str]],
           widths: Optional[List[float]] = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


def _exhibit(doc: Document, png: Optional[bytes], title: str, caption: str = ""):
    if png is None:
        return
    doc.add_heading(title, level=3)
    doc.add_picture(io.BytesIO(png), width=Inches(6.4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        _muted(doc, caption, size=8, italic=True)


def render_doc(path: Path, profile: CompanyProfile, kpi_set: KPISet,
               results: List[MetricResult], findings: List[Finding],
               images: Dict[str, bytes], specs: List, checks: List[str],
               anomaly_notes: List[str], period: str = "",
               tokens: Optional[Dict[str, str]] = None) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    # The token set rides on the document, matching `pdf.t` and `self.t` in the
    # other two renderers. Threading it through `_muted`'s seven call sites
    # would say the same thing seven more times.
    doc.t = dict(tokens or TOKENS["light"])
    _style(doc)
    cur = profile.identity.currency
    computed = [r for r in results if r.computed]

    # -- cover -----------------------------------------------------------
    _muted(doc, "PERFORMANCE REVIEW", size=10)
    doc.add_heading(profile.identity.name, level=0)
    _muted(doc, f"{profile.business_model.type.value.upper()} · "
                f"{profile.business_model.customer_type.value} · "
                f"{profile.identity.country} · {period}", size=11)
    _muted(doc, f"Prepared for the {profile.intent.audience.value} · primary "
                f"objective: {profile.intent.primary_objective.value.replace('_', ' ')} "
                f"· profile confidence {profile.confidence:.0%}")
    doc.add_paragraph()

    # -- 1 exec summary --------------------------------------------------
    doc.add_heading("1. Executive summary", level=1)
    doc.add_paragraph(
        "Every statement below is computed directly from the underlying data and "
        "reconciles to the scorecard that follows."
    )
    for f in findings[:5]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{SEVERITY_WORD.get(f.severity, '')} · {f.title}. ")
        run.bold = True
        p.add_run(f.statement)

    # -- 2 scorecard -----------------------------------------------------
    doc.add_heading("2. Scorecard", level=1)
    doc.add_paragraph(
        f"Grouped by Balanced Scorecard perspective; between two and seven KPIs "
        f"per perspective. Leading indicators are {kpi_set.leading_share:.0%} of the set."
    )
    for perspective in Perspective:
        group = [r for r in computed if r.kpi.perspective == perspective]
        if not group:
            continue
        doc.add_heading(perspective.value.replace("_", " ").title(), level=3)
        _table(doc,
               ["KPI", "Current", "12mo ago", "Target", "Cohort median", "Status"],
               [[r.kpi.short_name or r.kpi.name,
                 fmt_value(r.current, r.kpi.unit, cur),
                 fmt_value(r.prior_year, r.kpi.unit, cur),
                 fmt_value(r.target, r.kpi.unit, cur),
                 fmt_value(r.kpi.benchmark.p50, r.kpi.unit, cur) if r.kpi.benchmark else "—",
                 STATUS_WORD.get(r.status, "—")]
                for r in sorted(group, key=lambda x: int(x.kpi.tier))],
               widths=[2.0, 0.9, 0.9, 0.9, 1.0, 0.8])
        doc.add_paragraph()

    # -- 3 diagnostic ----------------------------------------------------
    doc.add_page_break()
    doc.add_heading("3. Diagnostic", level=1)
    bridge = next((f for f in findings if f.id == "arr_bridge"), None)
    if bridge:
        doc.add_paragraph(bridge.statement)
    _exhibit(doc, images.get("arr_bridge"), "ARR bridge, last 12 months",
             "Blue adds, red subtracts.")
    _exhibit(doc, images.get("cohort_heatmap"),
             "Revenue retention by acquisition cohort",
             "Rows below 100% at month 12 are cohorts where churn outran expansion.")

    # -- 4 deep dives ----------------------------------------------------
    doc.add_page_break()
    doc.add_heading("4. Deep dives", level=1)
    shown = {"arr_bridge", "cohort_heatmap"}
    finding_for = {}
    for f in findings:
        for kid in f.kpi_ids:
            finding_for.setdefault(kid, f)
    for spec in specs:
        if spec.id in shown or spec.id not in images:
            continue
        _exhibit(doc, images[spec.id], spec.title, spec.subtitle)
        related = next((f for f in findings if spec.id in f.id), None)
        if related is None:
            for kid, f in finding_for.items():
                if kid in spec.id or spec.id in kid:
                    related = f
                    break
        if related:
            p = doc.add_paragraph()
            p.add_run("Observation. ").bold = True
            p.add_run(related.statement)
            if related.recommendation:
                p2 = doc.add_paragraph()
                p2.add_run("Implication. ").bold = True
                p2.add_run(related.recommendation)

    # -- 5 benchmarks ----------------------------------------------------
    doc.add_page_break()
    doc.add_heading("5. Benchmarks", level=1)
    _exhibit(doc, images.get("benchmark_position"),
             "Position against the peer cohort median",
             "Direction normalised so positive always means better.")
    rows = [[r.kpi.short_name or r.kpi.name,
             fmt_value(r.current, r.kpi.unit, cur),
             fmt_value(r.kpi.benchmark.p25, r.kpi.unit, cur),
             fmt_value(r.kpi.benchmark.p50, r.kpi.unit, cur),
             fmt_value(r.kpi.benchmark.p75, r.kpi.unit, cur),
             (r.benchmark_position or "—").replace("_", " ")]
            for r in computed if r.kpi.benchmark]
    if rows:
        _table(doc, ["KPI", "You", "P25", "P50", "P75", "Position"], rows)
    _muted(doc, "Peer figures are illustrative placeholders assembled from public "
                "commentary, not a licensed benchmark dataset. Suitable for "
                "calibration, not for external reporting.", italic=True)

    # -- 6 risks ---------------------------------------------------------
    doc.add_heading("6. Risks and watch-list", level=1)
    risks = [f for f in findings if f.severity in ("critical", "high")]
    if not risks:
        doc.add_paragraph("No critical or high-severity issues were detected this period.")
    for f in risks:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{SEVERITY_WORD[f.severity]} · {f.title}. ").bold = True
        p.add_run(f.statement)

    # -- 7 actions -------------------------------------------------------
    doc.add_heading("7. Recommended actions", level=1)
    order = {"high": 0, "medium": 1, "low": 2, None: 3}
    owner_by_kpi = {r.kpi.id: r.kpi.owner_role for r in results}
    actionable = sorted([f for f in findings if f.recommendation],
                        key=lambda f: (order.get(f.impact, 3), order.get(f.effort, 3)))
    if actionable:
        _table(doc, ["#", "Action", "Moves", "Owner", "Impact", "Effort"],
               [[str(i), f.recommendation, ", ".join(f.kpi_ids[:2]) or "—",
                 next((owner_by_kpi.get(k) for k in f.kpi_ids if owner_by_kpi.get(k)), "—"),
                 (f.impact or "—").title(), (f.effort or "—").title()]
                for i, f in enumerate(actionable[:12], 1)],
               widths=[0.3, 2.6, 1.1, 0.7, 0.7, 0.6])
    else:
        doc.add_paragraph("No finding crossed an action threshold this period.")

    # -- 8 appendix ------------------------------------------------------
    doc.add_page_break()
    doc.add_heading("8. Appendix", level=1)

    doc.add_heading("Methodology", level=2)
    doc.add_paragraph(
        "KPIs were selected by evaluating each metric's applicability expression "
        "against this company profile, then scoring on objective alignment, "
        "Balanced Scorecard coverage, leading/lagging balance and audience fit. "
        "Every figure in this report is computed by deterministic code from the "
        "underlying fact tables; no number here was produced by a language model."
    )

    doc.add_heading("Assumptions and provenance", level=2)
    defaulted = [(p, s) for p, s in sorted(profile.provenance.items())
                 if s.startswith("benchmark_default")]
    if defaulted:
        for field, src in defaulted:
            doc.add_paragraph(f"{field} — {src}", style="List Bullet")
    else:
        doc.add_paragraph("No profile fields were filled from benchmark defaults.")

    doc.add_heading("Data quality", level=2)
    doc.add_paragraph(
        f"{len(checks)} reconciliation assertions passed before this report was "
        f"produced. The pipeline refuses to render on data that fails them."
    )
    for c in checks:
        doc.add_paragraph(c, style="List Bullet")

    if anomaly_notes:
        doc.add_heading("Known events in this dataset", level=2)
        for n in anomaly_notes:
            doc.add_paragraph(n, style="List Bullet")

    doc.add_heading("KPI definitions", level=2)
    for r in sorted(computed, key=lambda x: (x.kpi.perspective.value, int(x.kpi.tier))):
        k = r.kpi
        doc.add_heading(f"{k.name} ({k.perspective.value}, {k.timing.value})", level=3)
        doc.add_paragraph(f"Formula: {k.formula}")
        doc.add_paragraph(
            f"Owner: {k.owner_role} · Sources: {', '.join(k.source_systems) or 'n/a'} "
            f"· Frequency: {k.frequency}"
        )
        if k.interpretation:
            doc.add_paragraph(k.interpretation)
        if k.pitfalls:
            p = doc.add_paragraph()
            run = p.add_run(f"Pitfall: {k.pitfalls}")
            run.font.color.rgb = _rgb(doc.t["serious"])
        if k.benchmark:
            _muted(doc, f"Benchmark source: {k.benchmark.source}", size=8)

    if kpi_set.dropped:
        doc.add_heading("KPIs considered but not selected", level=2)
        for kid, reason in sorted(kpi_set.dropped.items()):
            doc.add_paragraph(f"{kid} — {reason}", style="List Bullet")

    doc.save(str(path))
    return path
